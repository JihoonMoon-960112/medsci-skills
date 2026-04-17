#!/usr/bin/env python3
"""Fill a Korean Word form template while preserving styles, tables, fonts, and page layout.

Core principles (DO NOT BREAK):
1. Always open existing template via Document(path) — never create from scratch.
2. Modify cell/paragraph TEXT only. Preserve all run-level styles.
3. Apply cantSplit to every row that gets filled (prevents page-break-mid-row).
4. Set Korean font with eastAsia attribute (run.font.name alone fails for Korean).
5. Validate: report empty cells and paragraphs that didn't match.
"""
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable

import yaml
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.table import _Cell


DEFAULT_KOREAN_FONT = "맑은 고딕"


# ---------- Style preservation helpers ----------

def _set_run_korean_font(run, font_name: str) -> None:
    """Set font for a run including eastAsia attribute (mandatory for Hangul)."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)


def _apply_cant_split(row) -> None:
    """Mark row to never split across pages."""
    trPr = row._tr.get_or_add_trPr()
    if trPr.find(qn("w:cantSplit")) is None:
        trPr.append(OxmlElement("w:cantSplit"))


def _replace_paragraph_text_keep_style(para: Paragraph, new_text: str,
                                        korean_font: str | None = None) -> None:
    """Replace the entire text content of a paragraph while keeping its style.

    Strategy: keep the first run's properties as the template style. Remove all
    other runs. Replace the first run's text with new_text. For multi-line
    content, split on \n and use w:br between lines (within same run-style block).
    """
    # Capture template run (first one) style by copying its rPr
    runs = para.runs
    template_rPr = None
    if runs:
        template_run_elem = runs[0]._element
        rPr = template_run_elem.find(qn("w:rPr"))
        if rPr is not None:
            template_rPr = rPr

    # Remove all existing runs
    for r in list(para._element.findall(qn("w:r"))):
        para._element.remove(r)

    # Add new run with the captured style
    new_run = OxmlElement("w:r")
    if template_rPr is not None:
        # Deep copy template rPr
        from copy import deepcopy
        new_run.append(deepcopy(template_rPr))

    # Split on \n — insert w:br between lines, w:t for text segments
    lines = new_text.split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            br = OxmlElement("w:br")
            new_run.append(br)
        if line:
            t = OxmlElement("w:t")
            t.text = line
            t.set(qn("xml:space"), "preserve")
            new_run.append(t)

    para._element.append(new_run)

    if korean_font:
        # Reapply Korean font to the new run
        from docx.text.run import Run
        run_obj = Run(new_run, para)
        _set_run_korean_font(run_obj, korean_font)


def _replace_cell_text(cell: _Cell, new_text: str,
                        korean_font: str | None = None) -> None:
    """Replace a cell's text content. Use the first paragraph as template."""
    if not cell.paragraphs:
        # Cell has no paragraph — add one
        cell.add_paragraph(new_text)
        if korean_font:
            for r in cell.paragraphs[0].runs:
                _set_run_korean_font(r, korean_font)
        return

    # Replace first paragraph, then remove the rest
    template_para = cell.paragraphs[0]

    # If the new content has multiple lines, we replace first paragraph
    # with the first line, and add additional paragraphs for remaining lines.
    lines = new_text.split("\n")

    _replace_paragraph_text_keep_style(template_para, lines[0],
                                        korean_font=korean_font)

    # Remove all paragraphs after the first
    for p in list(cell._tc.findall(qn("w:p")))[1:]:
        cell._tc.remove(p)

    # Add new paragraphs for remaining lines (cloning first paragraph's pPr)
    if len(lines) > 1:
        from copy import deepcopy
        first_p = cell._tc.find(qn("w:p"))
        first_pPr = first_p.find(qn("w:pPr")) if first_p is not None else None
        first_rPr = None
        first_r = first_p.find(qn("w:r")) if first_p is not None else None
        if first_r is not None:
            first_rPr = first_r.find(qn("w:rPr"))

        for line in lines[1:]:
            new_p = OxmlElement("w:p")
            if first_pPr is not None:
                new_p.append(deepcopy(first_pPr))
            new_r = OxmlElement("w:r")
            if first_rPr is not None:
                new_r.append(deepcopy(first_rPr))
            t = OxmlElement("w:t")
            t.text = line
            t.set(qn("xml:space"), "preserve")
            new_r.append(t)
            new_p.append(new_r)
            cell._tc.append(new_p)

        if korean_font:
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_korean_font(r, korean_font)


# ---------- FormFiller class ----------

@dataclass
class FillResult:
    matched: list[str] = field(default_factory=list)
    unmatched: list[str] = field(default_factory=list)


class FormFiller:
    def __init__(self, template_path: str | Path,
                 korean_font: str = DEFAULT_KOREAN_FONT,
                 blank_between_paragraphs: bool = True):
        self.path = Path(template_path).expanduser().resolve()
        if not self.path.exists():
            raise FileNotFoundError(self.path)
        self.doc = Document(str(self.path))
        self.korean_font = korean_font
        self.blank_between_paragraphs = blank_between_paragraphs
        self._filled_rows: set[int] = set()
        self._table_results = FillResult()
        self._paragraph_results = FillResult()

    # ---- Table cell filling ----

    def _cell_text(self, cell: _Cell) -> str:
        return "\n".join(p.text for p in cell.paragraphs).strip()

    def _label_match(self, cell_text: str, label: str) -> bool:
        # Normalize whitespace and newlines
        norm_cell = re.sub(r"\s+", "", cell_text)
        norm_label = re.sub(r"\s+", "", label)
        return norm_cell == norm_label

    def fill_table_kv(self, label: str, value: str) -> bool:
        """Find a cell whose text == label, fill the next cell on the right.

        Returns True if filled, False otherwise.
        Skips merged duplicate cells (same _tc reference).
        """
        for table in self.doc.tables:
            for row_idx, row in enumerate(table.rows):
                # Track unique cells in this row (skip merged duplicates)
                seen_tcs: set[int] = set()
                cells_in_row: list[_Cell] = []
                for c in row.cells:
                    if id(c._tc) not in seen_tcs:
                        seen_tcs.add(id(c._tc))
                        cells_in_row.append(c)

                for ci, cell in enumerate(cells_in_row):
                    if self._label_match(self._cell_text(cell), label):
                        # Found label cell. Fill the next cell on the right.
                        if ci + 1 < len(cells_in_row):
                            target = cells_in_row[ci + 1]
                            _replace_cell_text(target, value,
                                               korean_font=self.korean_font)
                            _apply_cant_split(row)
                            self._table_results.matched.append(label)
                            return True
        self._table_results.unmatched.append(label)
        return False

    # ---- Paragraph (section) filling ----

    def replace_paragraphs_after(self, header_text: str, new_content: str,
                                  stop_pattern: str | None = None) -> bool:
        """Find a paragraph matching header_text, then replace all paragraphs
        between this header and the next section header (or stop_pattern) with
        new_content.

        new_content is split by \n\n into separate paragraphs (preserving the
        style of the first replaced paragraph).
        """
        body = self.doc.element.body
        all_ps = list(self.doc.paragraphs)

        # Find header paragraph
        header_idx = None
        for i, p in enumerate(all_ps):
            if self._label_match(p.text, header_text):
                header_idx = i
                break

        if header_idx is None:
            self._paragraph_results.unmatched.append(header_text)
            return False

        # Determine end paragraph (next numbered section header or stop_pattern)
        if stop_pattern:
            end_re = re.compile(stop_pattern)
        else:
            # Match patterns like "1. ", "2. ", ... "18. "
            end_re = re.compile(r"^\s*\d+\.\s+\S")

        end_idx = len(all_ps)
        for i in range(header_idx + 1, len(all_ps)):
            if end_re.match(all_ps[i].text):
                end_idx = i
                break

        # Paragraphs to replace: header_idx+1 .. end_idx-1
        # Strategy: replace first paragraph in range, remove rest, add new paragraphs
        if header_idx + 1 >= end_idx:
            # No paragraphs between header and next section — just insert
            from copy import deepcopy
            template_p = all_ps[header_idx]._element
            template_pPr = template_p.find(qn("w:pPr"))
            template_r = template_p.find(qn("w:r"))
            template_rPr = template_r.find(qn("w:rPr")) if template_r is not None else None

            insert_after = template_p
            chunks = new_content.split("\n\n")
            for ci, chunk in enumerate(chunks):
                if ci > 0 and self.blank_between_paragraphs:
                    blank_p = OxmlElement("w:p")
                    insert_after.addnext(blank_p)
                    insert_after = blank_p
                new_p = OxmlElement("w:p")
                # New paragraph should NOT have header style — use default (no pPr)
                new_r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = chunk
                t.set(qn("xml:space"), "preserve")
                new_r.append(t)
                new_p.append(new_r)
                insert_after.addnext(new_p)
                # Apply Korean font
                from docx.text.run import Run
                _set_run_korean_font(Run(new_r, None), self.korean_font)
                insert_after = new_p
            self._paragraph_results.matched.append(header_text)
            return True

        # Replace first paragraph in range
        first_target = all_ps[header_idx + 1]
        chunks = new_content.split("\n\n")
        _replace_paragraph_text_keep_style(first_target, chunks[0],
                                           korean_font=self.korean_font)

        # Remove all paragraphs after first_target up to end_idx
        for i in range(header_idx + 2, end_idx):
            p_elem = all_ps[i]._element
            p_elem.getparent().remove(p_elem)

        # Add additional chunks as new paragraphs after first_target
        from copy import deepcopy
        first_target_elem = first_target._element
        first_pPr = first_target_elem.find(qn("w:pPr"))
        first_r = first_target_elem.find(qn("w:r"))
        first_rPr = first_r.find(qn("w:rPr")) if first_r is not None else None

        insert_after = first_target_elem
        for chunk in chunks[1:]:
            if self.blank_between_paragraphs:
                blank_p = OxmlElement("w:p")
                insert_after.addnext(blank_p)
                insert_after = blank_p
            new_p = OxmlElement("w:p")
            if first_pPr is not None:
                new_p.append(deepcopy(first_pPr))
            new_r = OxmlElement("w:r")
            if first_rPr is not None:
                new_r.append(deepcopy(first_rPr))
            t = OxmlElement("w:t")
            t.text = chunk
            t.set(qn("xml:space"), "preserve")
            new_r.append(t)
            new_p.append(new_r)
            insert_after.addnext(new_p)
            from docx.text.run import Run
            _set_run_korean_font(Run(new_r, None), self.korean_font)
            insert_after = new_p

        self._paragraph_results.matched.append(header_text)
        return True

    # ---- Single-paragraph in-place text replace ----

    def replace_paragraph_matching(self, matcher: str, new_text: str,
                                    mode: str = "startswith") -> bool:
        """Replace the entire text of the first paragraph that matches.

        mode: 'startswith' | 'contains' | 'exact'
        Preserves the paragraph's pPr and the first run's rPr (style).
        """
        for p in self.doc.paragraphs:
            text = p.text
            ok = False
            if mode == "startswith":
                ok = text.startswith(matcher)
            elif mode == "contains":
                ok = matcher in text
            elif mode == "exact":
                ok = text.strip() == matcher.strip()
            if ok:
                _replace_paragraph_text_keep_style(p, new_text,
                                                    korean_font=self.korean_font)
                self._paragraph_results.matched.append(f"<para>{matcher}")
                return True
        self._paragraph_results.unmatched.append(f"<para>{matcher}")
        return False

    # ---- Validation & save ----

    def validate(self) -> list[str]:
        warnings: list[str] = []
        for label in self._table_results.unmatched:
            warnings.append(f"[TABLE-MISS] Label not found: {label!r}")
        for header in self._paragraph_results.unmatched:
            warnings.append(f"[SECTION-MISS] Header not found: {header!r}")
        return warnings

    def report(self) -> str:
        n_table_ok = len(self._table_results.matched)
        n_table_miss = len(self._table_results.unmatched)
        n_para_ok = len(self._paragraph_results.matched)
        n_para_miss = len(self._paragraph_results.unmatched)
        return (
            f"Filled {n_table_ok} table cells, {n_para_ok} sections.\n"
            f"Missed: {n_table_miss} cells, {n_para_miss} sections."
        )

    def save(self, output_path: str | Path) -> Path:
        out = Path(output_path).expanduser().resolve()
        out.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(out))
        return out


# ---------- CLI ----------

def fill_from_yaml(template: Path, content_yaml: Path, output: Path) -> None:
    with open(content_yaml, "r", encoding="utf-8") as f:
        cfg = yaml.safe_load(f)

    protections = cfg.get("protections", {}) or {}
    korean_font = protections.get("korean_font", DEFAULT_KOREAN_FONT)
    blank_between = protections.get("blank_between_paragraphs", True)
    filler = FormFiller(template, korean_font=korean_font,
                         blank_between_paragraphs=blank_between)

    # Fill table key-value pairs
    for label, value in (cfg.get("table_kv") or {}).items():
        ok = filler.fill_table_kv(str(label), str(value))
        status = "OK " if ok else "MISS"
        print(f"  [{status}] table_kv: {label!r}")

    # Replace section content (between headers)
    for header, content in (cfg.get("section_replace") or {}).items():
        ok = filler.replace_paragraphs_after(str(header), str(content))
        status = "OK " if ok else "MISS"
        print(f"  [{status}] section: {header!r}")

    # Replace single paragraph in-place (e.g., title line)
    for matcher, content in (cfg.get("paragraph_replace") or {}).items():
        ok = filler.replace_paragraph_matching(str(matcher), str(content),
                                                mode="startswith")
        status = "OK " if ok else "MISS"
        print(f"  [{status}] paragraph: {matcher!r}")

    print()
    print(filler.report())
    print()

    warnings = filler.validate()
    for w in warnings:
        print(f"  WARN: {w}")

    saved = filler.save(output)
    print(f"\nSaved: {saved}")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", required=True, help="Path to template .docx")
    parser.add_argument("--content", required=True, help="Path to content YAML")
    parser.add_argument("--output", required=True, help="Output .docx path")
    args = parser.parse_args()

    fill_from_yaml(Path(args.template), Path(args.content), Path(args.output))


if __name__ == "__main__":
    main()
